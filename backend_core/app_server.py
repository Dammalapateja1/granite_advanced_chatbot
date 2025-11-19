
# backend_core/app_server.py

import threading
from pathlib import Path
from typing import List

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import (
    StreamingResponse,
    JSONResponse,
    HTMLResponse,
    Response,
)
from pydantic import BaseModel

from model_loader import get_llm, get_llm_tokenizer, load_models
from memory_handler import (
    add_message,
    format_history_for_prompt,
    clear_history,
    get_history,
)
from rag_engine import (
    add_document_from_file,
    query_corpus,
    corpus_size,
)
from utils_core import build_prompt, sanitize_stream_text

import torch
from transformers import TextIteratorStreamer

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "vector_store"
UPLOAD_DIR.mkdir(exist_ok=True, parents=True)

app = FastAPI(title="Granite Advanced Chatbot")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    session_id: str
    message: str
    use_rag: bool = True
    mode: str = "general"  # "general" | "coding" | "teacher" | "summarizer"


class SearchRequest(BaseModel):
    query: str
    top_k: int = 4


class ClearSessionRequest(BaseModel):
    session_id: str


class ExportRequest(BaseModel):
    session_id: str
    format: str = "txt"  # "txt" | "docx" | "pdf"


@app.on_event("startup")
def on_startup():
    print("[app_server] Starting up, loading models...")
    load_models()
    print("[app_server] Models ready.")


@app.get("/", response_class=HTMLResponse)
async def root():
    return """
    <html>
      <head><title>Granite Backend</title></head>
      <body style="font-family: sans-serif; background:#020617; color:#e5e7eb;">
        <h2>Granite backend is running ✅</h2>
        <p>Open <code>frontend_ui/index_ui.html</code> in your browser to use the chat UI.</p>
        <p>API docs: <a href="/docs">/docs</a> · Health: <a href="/health">/health</a></p>
      </body>
    </html>
    """


@app.post("/chat_stream")
async def chat_stream(req: ChatRequest):
    session_id = req.session_id or "default"
    user_message = req.message.strip()
    use_rag = req.use_rag
    mode = (req.mode or "general").lower()

    if not user_message:
        return JSONResponse({"error": "Empty message"}, status_code=400)

    def token_generator():
        history_prompt = format_history_for_prompt(session_id)
        rag_hits = query_corpus(user_message, top_k=4) if use_rag else []

        prompt = build_prompt(
            history_prompt=history_prompt,
            user_message=user_message,
            rag_context=rag_hits,
            mode=mode,
        )

        model = get_llm()
        tokenizer = get_llm_tokenizer()
        device = model.device if hasattr(model, "device") else torch.device("cpu")

        inputs = tokenizer(prompt, return_tensors="pt").to(device)
        streamer = TextIteratorStreamer(tokenizer, skip_special_tokens=True, skip_prompt=True)

        generate_kwargs = dict(
            **inputs,
            max_new_tokens=400,
            temperature=0.3,
            top_p=0.9,
            do_sample=True,
            streamer=streamer,
        )

        thread = threading.Thread(target=model.generate, kwargs=generate_kwargs)
        thread.start()

        collected = []
        for new_text in streamer:
            collected.append(new_text)
            yield sanitize_stream_text(new_text)

        full_answer = "".join(collected).strip()
        if full_answer:
            add_message(session_id, "user", user_message)
            add_message(session_id, "assistant", full_answer)

    return StreamingResponse(token_generator(), media_type="text/plain")


@app.post("/upload_file")
async def upload_file(
    file: UploadFile = File(...),
    source_name: str = Form(None),
):
    if source_name is None or not source_name.strip():
        source_name = file.filename or "uploaded_document"

    if not file.filename:
        return JSONResponse({"error": "No file name"}, status_code=400)

    save_path = UPLOAD_DIR / file.filename
    with open(save_path, "wb") as f:
        f.write(await file.read())

    try:
        chunks_added = add_document_from_file(str(save_path), source_name=source_name)
        total_chunks = corpus_size()
        return {
            "status": "ok",
            "file": file.filename,
            "chunks_added": chunks_added,
            "total_chunks": total_chunks,
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/search")
async def search_docs(req: SearchRequest):
    hits = query_corpus(req.query, top_k=req.top_k)
    return {"results": hits}


@app.post("/clear_session")
async def clear_session(req: ClearSessionRequest):
    clear_history(req.session_id)
    return {"status": "cleared", "session_id": req.session_id}


@app.post("/export_chat")
async def export_chat(req: ExportRequest):
    history = get_history(req.session_id)
    if not history:
        return JSONResponse({"error": "No messages for this session"}, status_code=404)

    fmt = req.format.lower()
    base_name = f"granite_chat_{req.session_id}"

    if fmt == "txt":
        lines = []
        for m in history:
            role = m["role"].capitalize()
            lines.append(f"{role}: {m['content']}")
            lines.append("")
        text = "\n".join(lines)
        return Response(
            text,
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{base_name}.txt"'},
        )

    elif fmt == "docx":
        from io import BytesIO
        from docx import Document
        buffer = BytesIO()
        doc = Document()
        doc.add_heading("Granite Chat Export", level=1)
        doc.add_paragraph(f"Session ID: {req.session_id}")
        doc.add_paragraph("")
        for m in history:
            p = doc.add_paragraph()
            role_run = p.add_run(f"{m['role'].capitalize()}: ")
            role_run.bold = True
            p.add_run(m["content"])
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base_name}.docx"'},
        )

    elif fmt == "pdf":
        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        import textwrap

        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        margin = 1 * inch
        y = height - margin

        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, "Granite Chat Export")
        y -= 18

        c.setFont("Helvetica", 10)
        c.drawString(margin, y, f"Session ID: {req.session_id}")
        y -= 24

        wrap_width = 95

        for m in history:
            role = m["role"].capitalize()
            content = (m["content"] or "").replace("\r", "")
            if y <= margin:
                c.showPage()
                y = height - margin
            c.setFont("Helvetica-Bold", 10)
            c.drawString(margin, y, f"{role}:")
            y -= 14
            c.setFont("Helvetica", 10)
            for line in content.split("\n"):
                for subline in textwrap.wrap(line, width=wrap_width):
                    if y <= margin:
                        c.showPage()
                        y = height - margin
                        c.setFont("Helvetica", 10)
                    c.drawString(margin + 20, y, subline)
                    y -= 12
            y -= 8

        c.showPage()
        c.save()
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{base_name}.pdf"'},
        )

    else:
        return JSONResponse({"error": "Unsupported format"}, status_code=400)


@app.get("/health")
async def health():
    return {"status": "ok", "corpus_chunks": corpus_size()}
